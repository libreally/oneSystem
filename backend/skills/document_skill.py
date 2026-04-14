"""
文档处理 Skill
支持公文转换、格式调整等功能
"""
from .base_skill import BaseSkill
from typing import Dict, Any
import os
import logging

logger = logging.getLogger(__name__)


class DocumentSkill(BaseSkill):
    """文档处理 Skill"""
    
    def __init__(self):
        super().__init__(
            skill_id='doc_processor',
            name='文档处理',
            description='支持公文转换、格式调整、文档生成等功能'
        )
    
    def validate_params(self, params: Dict[str, Any]) -> tuple:
        """验证参数"""
        errors = []
        
        if 'action' not in params:
            errors.append('缺少 action 参数')
        
        if params.get('action') == 'convert_to_official':
            if 'file_path' not in params and 'content' not in params:
                errors.append('需要提供 file_path 或 content')
            if 'doc_type' not in params:
                errors.append('缺少 doc_type 参数（通知/报告/请示等）')
        
        return len(errors) == 0, errors
    
    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """执行文档处理"""
        try:
            action = params.get('action')
            
            if action == 'convert_to_official':
                return self._convert_to_official(params)
            elif action == 'format_adjust':
                return self._format_adjust(params)
            elif action == 'generate_doc':
                return self._generate_doc(params)
            else:
                return {
                    'success': False,
                    'message': f'未知的操作类型：{action}'
                }
                
        except Exception as e:
            logger.error(f"文档处理失败：{str(e)}")
            return {
                'success': False,
                'message': f'处理失败：{str(e)}'
            }
    
    def _convert_to_official(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """转换为标准公文格式"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from backend.services.config_service import config_service
        
        content = params.get('content', '')
        file_path = params.get('file_path')
        doc_type = params.get('doc_type', '通知')
        output_path = params.get('output_path', 'output.docx')
        
        # 从配置中心加载公文模板配置
        template_config = config_service.get_config('document_template', 'standard')
        
        # 如果提供了文件路径，读取文件内容
        if file_path and os.path.exists(file_path):
            # TODO: 读取现有文档内容
            pass
        
        # 创建新文档
        doc = Document()
        
        # 设置标题
        title = params.get('title', f'关于 XXX 的{doc_type}')
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 应用配置中的格式
        if template_config:
            # 设置字体和字号
            if 'font' in template_config and 'font_size' in template_config:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = template_config['font']
                        run.font.size = Pt(template_config['font_size'])
        
        # 添加正文
        if content:
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    # 应用配置中的行间距和首行缩进
                    if template_config:
                        if 'line_spacing' in template_config:
                            p.paragraph_format.line_spacing = template_config['line_spacing']
                        if 'margin' in template_config and 'left' in template_config['margin']:
                            p.paragraph_format.first_line_indent = Cm(template_config['margin']['left'] * 0.29)
                    else:
                        # 默认值
                        p.paragraph_format.line_spacing = Cm(0.75)
                        p.paragraph_format.first_line_indent = Cm(0.74)
        
        # 保存文档
        doc.save(output_path)
        
        self.log_execution('convert_to_official', f'已生成公文：{output_path}')
        
        return {
            'success': True,
            'message': '公文转换成功',
            'file_path': os.path.abspath(output_path),
            'data': {
                'doc_type': doc_type,
                'output_path': output_path,
                'template_used': template_config is not None
            }
        }
    
    def _format_adjust(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """调整文档格式"""
        # TODO: 实现格式调整功能
        return {
            'success': True,
            'message': '格式调整功能开发中'
        }
    
    def _generate_doc(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """生成文档"""
        from docx import Document
        
        template = params.get('template', 'default')
        data = params.get('data', {})
        output_path = params.get('output_path', 'generated.docx')
        
        doc = Document()
        
        # 根据模板和数据生成内容
        title = data.get('title', '生成的文档')
        doc.add_heading(title, level=1)
        
        content = data.get('content', '')
        if content:
            doc.add_paragraph(content)
        
        doc.save(output_path)
        
        return {
            'success': True,
            'message': '文档生成成功',
            'file_path': os.path.abspath(output_path)
        }
